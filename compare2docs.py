from docx import Document
from colorama import Fore, Style

def compare_word_documents(doc_path1, doc_path2, output_path):
    doc1 = Document(doc_path1)
    doc2 = Document(doc_path2)
    
    differences = []
    
    for para1, para2 in zip(doc1.paragraphs, doc2.paragraphs):
        if para1.text != para2.text:
            differences.append(f"{Fore.RED}Left: {para1.text}\n{Fore.GREEN}Right: {para2.text}\n{Style.RESET_ALL}")
    
    with open(output_path, 'w', encoding='utf-8') as output_file:
        output_file.write("\n".join(differences))

# 要比較的兩個Word文檔路徑
doc_path1 = r'112年宜蘭縣下水道營運管理系統精進及維護計畫_期中報告書_1130102_V6(無標記).docx'
doc_path2 = r'112年宜蘭縣下水道營運管理系統精進及維護計畫_期末報告書_1130328_V18.docx'

# 輸出差異的.txt文件路徑
output_path = 'differences.txt'

# 執行文檔比較並將差異列出到.txt文件中
compare_word_documents(doc_path1, doc_path2, output_path)
