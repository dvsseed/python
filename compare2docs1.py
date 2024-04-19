from docx import Document
# from colorama import Fore, Style
import difflib

def compare_word_documents(doc_path1, doc_path2, output_path):
    doc1 = Document(doc_path1)
    doc2 = Document(doc_path2)
    
    differences = []
    
    for para1, para2 in zip(doc1.paragraphs, doc2.paragraphs):
        if para1.text != para2.text:
            diff = difflib.HtmlDiff().make_file([para1.text], [para2.text], para1.text, para2.text, context=True)
            differences.append(diff)
    
    with open(output_path, 'w', encoding='utf-8') as output_file:
        output_file.write("<html><head><style>td { font-family: monospace; }</style></head><body>")
        for diff in differences:
            output_file.write(diff)
        output_file.write("</body></html>")

def compare_word_documents_x(doc_path1, doc_path2, output_path):
    doc1 = Document(doc_path1)
    doc2 = Document(doc_path2)
    
    differences = []
    
    for para1, para2 in zip(doc1.paragraphs, doc2.paragraphs):
        if para1.text != para2.text:
            differences.append(f"{Fore.RED}Left: {para1.text}{Style.RESET_ALL}\n{Fore.GREEN}Right: {para2.text}{Style.RESET_ALL}\n")
        else:
            differences.append(f"{para1.text}\n")
    
    with open(output_path, 'w', encoding='utf-8') as output_file:
        output_file.write("\n".join(differences))

# 要比较的两个Word文档路径
doc_path1 = r'C:\Users\dvsse\112年宜蘭縣下水道營運管理系統精進及維護計畫_期中報告書_1130102_V6(無標記).docx'
doc_path2 = r'C:\Users\dvsse\112年宜蘭縣下水道營運管理系統精進及維護計畫_期末報告書_1130328_V18.docx'

# 输出差异的.txt文件路径
# output_path = r'C:\Users\dvsse\differences.txt'
# 输出差异的HTML文件路径
output_path = r'C:\Users\dvsse\differences.html'

# 执行文档比较并将差异列出到.txt文件中
compare_word_documents(doc_path1, doc_path2, output_path)
